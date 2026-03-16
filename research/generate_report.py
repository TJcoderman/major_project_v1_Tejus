#!/usr/bin/env python3
"""
Generate Graphic Era University Project Progress Report (DOCX)
Matches the exact format from the phase1_report template.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "SecureBank_Project_Report.docx")
FIGURES_DIR = os.path.join(SCRIPT_DIR, "figures")


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_paragraph(doc, text, font_name="Bookman Old Style", size=12,
                           bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                           space_after=6, space_before=0, first_line_indent=None):
    """Add a formatted paragraph."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return p


def add_heading_styled(doc, text, level=1):
    """Add heading with Bookman Old Style."""
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.name = "Bookman Old Style"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_page_break(doc):
    doc.add_page_break()


def create_title_page(doc):
    """Page 1: Title Page."""
    # Title header
    add_formatted_paragraph(doc, "Project Progress Report on",
                           size=16, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                           space_after=4, space_before=72)

    # Horizontal line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("═" * 60)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Project title
    add_formatted_paragraph(doc,
        "Enhancing Mobile Banking Security through Behavior-Based Continuous Authentication",
        size=18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, space_before=8)

    # Horizontal line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("═" * 60)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Submitted in partial fulfillment
    add_formatted_paragraph(doc,
        "Submitted in partial fulfilment of the requirement for the award of the degree of",
        size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=24)

    add_formatted_paragraph(doc, "BACHELOR OF TECHNOLOGY",
        size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

    add_formatted_paragraph(doc, "IN",
        size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

    add_formatted_paragraph(doc, "COMPUTER SCIENCE & ENGINEERING",
        size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

    add_formatted_paragraph(doc, "Submitted by:",
        size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    # Student table
    students = [
        ("Bela Diwan", "2021947"),
        ("Gautam Nautiyal", "2021924"),
        ("Tejus Kapoor", "2021985"),
        ("Tanush Malhotra", "2021984"),
    ]

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (name, roll) in enumerate(students):
        for j, val in enumerate([name, roll]):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 1 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.name = "Bookman Old Style"
            run.font.size = Pt(12)
            run.font.bold = True

    # Guide
    add_formatted_paragraph(doc, "", size=8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Under the Guidance of")
    run.font.name = "Bookman Old Style"
    run.font.size = Pt(11)
    run.font.italic = True

    add_formatted_paragraph(doc, "Ms. Anamika Sharma",
        size=13, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    add_formatted_paragraph(doc, "Professor",
        size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    add_formatted_paragraph(doc, "Project Team ID:  MP2025IOT6",
        size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # University name
    add_formatted_paragraph(doc, "", size=24)
    add_formatted_paragraph(doc, "Department of Computer Science and Engineering",
        size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12)

    add_formatted_paragraph(doc, "", size=8)

    add_formatted_paragraph(doc, "Graphic Era (Deemed to be University)",
        size=14, bold=True, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    add_formatted_paragraph(doc, "Dehradun, Uttarakhand",
        size=13, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    add_formatted_paragraph(doc, "February-2026",
        size=13, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)


def create_declaration_page(doc):
    """Page 2: Candidate's Declaration."""
    add_page_break(doc)

    add_formatted_paragraph(doc, "CANDIDATE'S DECLARATION",
        size=16, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18, space_before=36)

    # Declaration text
    add_formatted_paragraph(doc,
        'We hereby certify that the work which is being presented in the Synopsis entitled '
        '"Enhancing Mobile Banking Security through Behavior-Based Continuous Authentication" '
        'in partial fulfillment of the requirements for the award of the Degree of Bachelor of Technology '
        'in Computer Science and Engineering in the Department of Computer Science and Engineering '
        'of the Graphic Era (Deemed to be University), Dehradun shall be carried out by the '
        'undersigned under the supervision of Ms. Anamika Sharma, Professor, Department of '
        'Computer Science and Engineering, Graphic Era (Deemed to be University), Dehradun.',
        size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12,
        first_line_indent=1.27)

    # Student list
    students = [
        ("Bela Diwan", "2021947"),
        ("Gautam Nautiyal", "2021924"),
        ("Tejus Kapoor", "2021985"),
        ("Tanush Malhotra", "2021984"),
    ]

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (name, roll) in enumerate(students):
        for j, val in enumerate([name, roll]):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if j == 1 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.name = "Bookman Old Style"
            run.font.size = Pt(12)
            run.font.bold = True

    add_formatted_paragraph(doc, "", size=12)
    add_formatted_paragraph(doc,
        'The above mentioned students shall be working under the supervision of the undersigned on '
        'the "Enhancing Mobile Banking Security through Behavior-Based Continuous Authentication"',
        size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=24,
        first_line_indent=1.27)

    # Signatures
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, (line1, line2) in enumerate([
        ("Signature", "Supervisor"),
        ("Signature", "Head of the Department")
    ]):
        cell = table.cell(0, j)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line1)
        run.font.name = "Bookman Old Style"
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.italic = True

        cell = table.cell(1, j)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line2)
        run.font.name = "Bookman Old Style"
        run.font.size = Pt(12)
        run.font.bold = True


def create_toc(doc):
    """Page 3: Table of Contents."""
    add_page_break(doc)

    add_formatted_paragraph(doc, "Table of Contents",
        size=18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24, space_before=36)

    chapters = [
        ("Chapter 1", "Introduction and Problem Statement", "1-3"),
        ("Chapter 2", "Objectives", "4"),
        ("Chapter 3", "Project Work Carried Out", "5-8"),
        ("Chapter 4", "Results and Analysis", "9-11"),
        ("Chapter 5", "Future Work Plan", "12"),
        ("Chapter 6", "Weekly Task", "13"),
        ("", "References", "14"),
    ]

    table = doc.add_table(rows=len(chapters) + 1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, header in enumerate(["Chapter No.", "Description", "Page No."]):
        cell = table.cell(0, j)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.name = "Bookman Old Style"
        run.font.size = Pt(12)
        run.font.bold = True

    for i, (ch, desc, pg) in enumerate(chapters):
        for j, val in enumerate([ch, desc, pg]):
            cell = table.cell(i + 1, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j != 1 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.name = "Bookman Old Style"
            run.font.size = Pt(12)


def create_chapter1(doc):
    """Chapter 1: Introduction and Problem Statement."""
    add_page_break(doc)

    add_heading_styled(doc, "Chapter 1: Introduction and Problem Statement", level=1)

    add_heading_styled(doc, "1.1 Introduction", level=2)

    add_formatted_paragraph(doc,
        "The rapid proliferation of mobile banking has fundamentally transformed how individuals interact "
        "with financial services. With over 2.5 billion users worldwide accessing banking services through "
        "smartphones, the convenience and accessibility of mobile banking have become indispensable. However, "
        "this widespread adoption has simultaneously introduced significant security challenges. Traditional "
        "authentication mechanisms—PINs, passwords, and even fingerprint recognition—provide only a single "
        "point of verification at the moment of login. Once a user is authenticated, the session remains open "
        "and vulnerable to unauthorized access through shoulder surfing, device theft, session hijacking, or "
        "social engineering attacks.",
        first_line_indent=1.27)

    add_formatted_paragraph(doc,
        "Behavioral biometrics represents a paradigm shift in mobile security by enabling continuous, "
        "transparent authentication. Unlike physiological biometrics (such as fingerprints or facial recognition), "
        "behavioral biometrics analyzes how a user interacts with their device—their unique typing rhythm, "
        "touch pressure, swipe velocity, and device handling patterns. This approach offers several key "
        "advantages: it operates transparently without disrupting the user experience, it provides continuous "
        "authentication throughout the entire session rather than at a single point, and it is significantly "
        "harder to replicate compared to static credentials.",
        first_line_indent=1.27)

    add_heading_styled(doc, "1.2 Problem Statement", level=2)

    add_formatted_paragraph(doc,
        "Current mobile banking security suffers from three critical limitations:",
        first_line_indent=1.27)

    problems = [
        "Single-Point Authentication: Traditional PIN/password systems verify identity only at login. "
        "If a device is compromised after authentication (e.g., the user leaves the phone unlocked), "
        "the banking session remains accessible to any user.",

        "Static Credential Vulnerability: PINs and passwords can be observed, guessed, or "
        "stolen through phishing attacks, keyloggers, or social engineering. A 2024 report by "
        "Juniper Research estimates that mobile banking fraud losses will exceed $10 billion annually by 2027.",

        "User Experience Trade-off: Stronger security measures (complex passwords, frequent "
        "re-authentication) degrade the user experience, leading to user frustration and potential "
        "abandonment of security practices."
    ]

    for i, problem in enumerate(problems, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(1.27)
        run = p.add_run(f"{i}. {problem}")
        run.font.name = "Bookman Old Style"
        run.font.size = Pt(12)

    add_formatted_paragraph(doc,
        "This project addresses these challenges by developing SecureBank, a multi-modal behavioral "
        "biometric authentication system that continuously verifies user identity during mobile banking "
        "interactions using keystroke dynamics, touch behavior patterns, and motion sensor data—all captured "
        "transparently without any additional user effort.",
        first_line_indent=1.27, space_before=6)

    add_heading_styled(doc, "1.3 Scope of the Project", level=2)

    add_formatted_paragraph(doc,
        "The project encompasses the development of a fully functional Android banking application "
        "with integrated behavioral biometric data collection, machine learning-based classification, "
        "and comprehensive evaluation across diverse user demographics. The system captures three "
        "modalities of behavioral data: (i) PIN keystroke timing (dwell time, flight time, touch coordinates), "
        "(ii) touch interaction patterns (swipe velocity, acceleration, tap precision), and "
        "(iii) motion sensor readings (accelerometer, gyroscope, device orientation). These multi-modal "
        "features are processed through an enrollment-relative deviation approach and classified using "
        "Random Forest, MLP Neural Network, and One-Class SVM algorithms.",
        first_line_indent=1.27)


def create_chapter2(doc):
    """Chapter 2: Objectives."""
    add_page_break(doc)

    add_heading_styled(doc, "Chapter 2: Objectives", level=1)

    add_formatted_paragraph(doc,
        "The primary objectives of this project are as follows:",
        space_after=12)

    objectives = [
        "To design and develop a functional Android mobile banking application (SecureBank) "
        "with integrated, non-intrusive behavioral biometric data collection capabilities.",

        "To implement transparent data collectors for three behavioral modalities: keystroke dynamics "
        "(PIN entry timing), touch behavior (swipe gestures, tap patterns), and motion sensor data "
        "(accelerometer and gyroscope readings).",

        "To develop a novel enrollment-relative deviation feature engineering approach that computes "
        "behavioral features as deviations from each user's enrolled baseline profile.",

        "To extract and analyze 124 behavioral features from multi-modal interaction data for "
        "user authentication.",

        "To train and evaluate three machine learning classifiers—Random Forest, Multi-Layer Perceptron "
        "(MLP), and One-Class SVM—for genuine vs. impostor classification.",

        "To achieve authentication accuracy exceeding 90% with an Equal Error Rate (EER) below 10% "
        "using the best-performing classifier.",

        "To evaluate the system's performance across diverse age demographics (18-55 years) and analyze "
        "age-related variations in behavioral biometric patterns.",

        "To generate publication-quality visualizations including ROC curves, confusion matrices, "
        "t-SNE clustering plots, and feature importance rankings for research dissemination.",
    ]

    for i, obj in enumerate(objectives, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.hanging_indent = Cm(0.76)
        run = p.add_run(f"{i}. {obj}")
        run.font.name = "Bookman Old Style"
        run.font.size = Pt(12)


def create_chapter3(doc):
    """Chapter 3: Project Work Carried Out."""
    add_page_break(doc)

    add_heading_styled(doc, "Chapter 3: Project Work Carried Out", level=1)

    add_heading_styled(doc, "3.1 System Architecture", level=2)

    add_formatted_paragraph(doc,
        "The SecureBank system is implemented as a native Android application using Kotlin with "
        "Jetpack Compose for the user interface and Hilt for dependency injection. The architecture "
        "follows a layered design comprising four principal components:",
        first_line_indent=1.27)

    layers = [
        ("Banking Application Layer", "Implements core banking functionality including user login "
         "via 6-digit PIN, account dashboard with balance display, and fund transfer operations. "
         "A custom numeric keypad captures precise touch coordinates and timing for each PIN digit."),

        ("Data Collection Layer", "Three transparent collectors operate simultaneously: "
         "(a) Keystroke Dynamics Collector captures key-press/release times, dwell time, flight time, "
         "and touch coordinates; (b) Touch Behavior Collector records swipe direction, velocity, "
         "acceleration, and duration for all screen interactions; (c) Motion Sensor Collector "
         "samples accelerometer and gyroscope at ~500 Hz, capturing device orientation and handling patterns."),

        ("Feature Extraction Layer", "Extracts 124 behavioral features using enrollment-relative "
         "deviation computation. For each feature f, absolute deviation |f_session - f_enrollment| "
         "and relative deviation |f_session - f_enrollment| / |f_enrollment| are computed."),

        ("ML Classification Layer", "Three classifiers evaluate authentication decisions: "
         "Random Forest (300 trees), MLP Neural Network (128→64→32 architecture), "
         "and One-Class SVM (RBF kernel, anomaly detection approach)."),
    ]

    for name, desc in layers:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(1.27)
        run = p.add_run(f"• {name}: ")
        run.font.name = "Bookman Old Style"
        run.font.size = Pt(12)
        run.font.bold = True
        run = p.add_run(desc)
        run.font.name = "Bookman Old Style"
        run.font.size = Pt(12)

    # Architecture diagram
    arch_path = os.path.join(FIGURES_DIR, "system_architecture.png")
    if not os.path.exists(arch_path):
        arch_path = os.path.join(SCRIPT_DIR, "paper", "figures", "system_architecture.png")

    if os.path.exists(arch_path):
        add_formatted_paragraph(doc, "", size=6)
        doc.add_picture(arch_path, width=Inches(4.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_formatted_paragraph(doc, "Figure 1: System Architecture of SecureBank",
            size=10, bold=True, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    add_heading_styled(doc, "3.2 Data Collection Protocol", level=2)

    add_formatted_paragraph(doc,
        "We collected behavioral data from 40 participants (8 real participants, 32 statistically "
        "augmented) across four age demographics: 18-25 years (14 participants), 26-35 years "
        "(14 participants), 36-45 years (7 participants), and 46-55 years (5 participants). Each "
        "participant completed three types of experimental sessions:",
        first_line_indent=1.27)

    sessions = [
        ("Enrollment Session", "Five complete attempts of the banking workflow (PIN entry, text "
         "typing, touch interaction, free browsing) to establish the user's behavioral baseline."),
        ("Genuine Session", "Three attempts by the same user to simulate normal authentication."),
        ("Impostor Session", "Three attempts by a different user to simulate unauthorized access."),
    ]

    for name, desc in sessions:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(1.27)
        run = p.add_run(f"• {name}: ")
        run.font.name = "Bookman Old Style"
        run.font.size = Pt(12)
        run.font.bold = True
        run = p.add_run(desc)
        run.font.name = "Bookman Old Style"
        run.font.size = Pt(12)

    add_heading_styled(doc, "3.3 Feature Engineering", level=2)

    add_formatted_paragraph(doc,
        "A total of 124 enrollment-relative deviation features are extracted across three "
        "behavioral modalities:",
        first_line_indent=1.27)

    features = [
        ("PIN Keystroke Features (24 raw features)", "Mean, standard deviation, median, Q25, Q75, "
         "skewness, kurtosis, and IQR of dwell times and flight times. Per-digit-position rhythm "
         "features capture the typing cadence at each position in the 6-digit PIN sequence."),
        ("Touch Interaction Features (9 raw features)", "Mean and standard deviation of swipe duration, "
         "velocity, and acceleration; maximum velocity; tap-to-swipe ratio; and total event count."),
        ("Motion Sensor Features (7 raw features)", "Mean and standard deviation of accelerometer "
         "and gyroscope magnitudes, and individual axis means."),
    ]

    for name, desc in features:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(1.27)
        run = p.add_run(f"• {name}: ")
        run.font.name = "Bookman Old Style"
        run.font.size = Pt(12)
        run.font.bold = True
        run = p.add_run(desc)
        run.font.name = "Bookman Old Style"
        run.font.size = Pt(12)

    add_formatted_paragraph(doc,
        "Each raw feature is transformed into absolute deviation, relative deviation, and the raw "
        "session value, yielding 40 × 3 = 120 features plus 4 aggregate deviation statistics, "
        "totaling 124 features per session.",
        first_line_indent=1.27, space_before=6)

    add_heading_styled(doc, "3.4 Technology Stack", level=2)

    tech = [
        ("Platform", "Android (API 26+, Kotlin 2.2.10)"),
        ("UI Framework", "Jetpack Compose with Material Design 3"),
        ("Dependency Injection", "Dagger Hilt 2.56"),
        ("Database", "Room Persistence Library 2.7.1"),
        ("Build System", "Gradle 9.2.1 with KSP 2.2.10-2.0.2"),
        ("ML Framework", "Scikit-learn 1.3.0 (Python)"),
        ("Visualization", "Matplotlib 3.7.0, Seaborn 0.12.0"),
    ]

    table = doc.add_table(rows=len(tech) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header
    for j, h in enumerate(["Component", "Technology / Version"]):
        cell = table.cell(0, j)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = "Bookman Old Style"
        run.font.size = Pt(11)
        run.font.bold = True

    for i, (comp, tech_val) in enumerate(tech):
        for j, val in enumerate([comp, tech_val]):
            cell = table.cell(i + 1, j)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = "Bookman Old Style"
            run.font.size = Pt(11)

    add_formatted_paragraph(doc, "Table 1: Technology Stack",
        size=10, bold=True, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=4)


def create_chapter4(doc):
    """Chapter 4: Results and Analysis."""
    add_page_break(doc)

    add_heading_styled(doc, "Chapter 4: Results and Analysis", level=1)

    add_heading_styled(doc, "4.1 Model Performance Comparison", level=2)

    add_formatted_paragraph(doc,
        "Table 2 presents the comparative performance of all three classifiers evaluated using "
        "5-fold stratified cross-validation on the complete dataset of 80 sessions (40 genuine, "
        "40 impostor) from 40 participants.",
        first_line_indent=1.27)

    # Results table
    table = doc.add_table(rows=4, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = ["Model", "Accuracy", "Precision", "Recall", "F1 Score", "AUC-ROC", "EER"]
    results = [
        ["Random Forest", "0.9250", "0.9250", "0.9250", "0.9250", "0.9812", "0.0750"],
        ["MLP Neural Net", "0.8250", "0.7826", "0.9000", "0.8372", "0.9363", "0.1375"],
        ["One-Class SVM", "0.8125", "0.9032", "0.7000", "0.7887", "0.9344", "0.0625"],
    ]

    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = "Bookman Old Style"
        run.font.size = Pt(10)
        run.font.bold = True

    for i, row in enumerate(results):
        for j, val in enumerate(row):
            cell = table.cell(i + 1, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.name = "Bookman Old Style"
            run.font.size = Pt(10)
            if i == 0:
                run.font.bold = True

    add_formatted_paragraph(doc, "Table 2: Model Performance Comparison",
        size=10, bold=True, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=4)

    add_formatted_paragraph(doc,
        "The Random Forest classifier achieves the highest overall accuracy (92.50%), F1 score "
        "(0.925), and AUC-ROC (0.981). The One-Class SVM achieves the lowest EER (6.25%), "
        "making it particularly suitable for anomaly-based deployment scenarios where impostor "
        "training data may not be available.",
        first_line_indent=1.27, space_before=8)

    add_heading_styled(doc, "4.2 ROC Curves", level=2)

    roc_path = os.path.join(FIGURES_DIR, "fig3_roc_curves.png")
    if os.path.exists(roc_path):
        doc.add_picture(roc_path, width=Inches(4.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_formatted_paragraph(doc, "Figure 2: ROC Curves for All Three Models (RF AUC=0.981)",
            size=10, bold=True, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading_styled(doc, "4.3 Confusion Matrices", level=2)

    cm_path = os.path.join(FIGURES_DIR, "fig4_confusion_matrices.png")
    if os.path.exists(cm_path):
        doc.add_picture(cm_path, width=Inches(5.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_formatted_paragraph(doc, "Figure 3: Confusion Matrices for All Three Classifiers",
            size=10, bold=True, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading_styled(doc, "4.4 Feature Importance", level=2)

    add_formatted_paragraph(doc,
        "The Random Forest feature importance analysis reveals that PIN keystroke features "
        "dominate the classification decision, contributing over 60% of total importance. "
        "The top three features are: (1) PIN dwell time deviation (6.75%), (2) touch size "
        "deviation (5.73%), and (3) touch Y-coordinate deviation (5.56%).",
        first_line_indent=1.27)

    fi_path = os.path.join(FIGURES_DIR, "fig7_feature_importance.png")
    if os.path.exists(fi_path):
        doc.add_picture(fi_path, width=Inches(4.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_formatted_paragraph(doc, "Figure 4: Top 20 Feature Importance (Random Forest)",
            size=10, bold=True, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading_styled(doc, "4.5 Age-Group Analysis", level=2)

    add_formatted_paragraph(doc,
        "Analysis across four age demographics reveals systematic behavioral variations: dwell time "
        "increases with age (18-25: fastest, 46-55: ~45% slower), touch velocity decreases with "
        "age, and PIN entry time shows the most significant variation. Importantly, the enrollment-relative "
        "approach normalizes these differences, making the system inherently age-adaptive.",
        first_line_indent=1.27)

    age_path = os.path.join(FIGURES_DIR, "fig6_age_group_analysis.png")
    if os.path.exists(age_path):
        doc.add_picture(age_path, width=Inches(4.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_formatted_paragraph(doc, "Figure 5: Behavioral Variations Across Age Demographics",
            size=10, bold=True, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading_styled(doc, "4.6 Model Performance Comparison Chart", level=2)

    mc_path = os.path.join(FIGURES_DIR, "fig8_model_comparison.png")
    if os.path.exists(mc_path):
        doc.add_picture(mc_path, width=Inches(5.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_formatted_paragraph(doc, "Figure 6: Comparative Performance Across All Metrics",
            size=10, bold=True, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)


def create_chapter5(doc):
    """Chapter 5: Future Work Plan."""
    add_page_break(doc)

    add_heading_styled(doc, "Chapter 5: Future Work Plan", level=1)

    add_formatted_paragraph(doc,
        "The following activities are planned for the next phase of the project:",
        space_after=12)

    future_work = [
        ("Real-Time On-Device Inference",
         "Deploy the trained Random Forest model for real-time authentication on Android using "
         "TensorFlow Lite or ONNX Runtime. This will enable continuous authentication during "
         "live banking sessions without requiring server-side inference."),

        ("Expanded Dataset Collection",
         "Collect behavioral data from 100+ real participants across diverse demographics, "
         "device types, and environmental conditions (walking, seated, outdoor) to validate "
         "system robustness in real-world scenarios."),

        ("Adversarial Attack Evaluation",
         "Evaluate system robustness against active adversarial attacks including mimicry attacks "
         "(where an impostor attempts to replicate the genuine user's typing pattern) and "
         "replay attacks."),

        ("Deep Learning Exploration",
         "Implement and evaluate deep learning approaches such as Siamese Neural Networks for "
         "one-shot behavioral verification and LSTM-based models for temporal behavior modeling."),

        ("Adaptive Enrollment",
         "Implement adaptive enrollment mechanisms that continuously update the user's behavioral "
         "profile to account for natural behavioral drift over time (e.g., typing speed changes "
         "due to fatigue, device familiarity, or physical conditions)."),

        ("Integration with Banking APIs",
         "Connect the behavioral authentication layer with actual banking APIs and risk scoring "
         "systems to demonstrate end-to-end deployment viability in production environments."),
    ]

    for i, (title, desc) in enumerate(future_work, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.hanging_indent = Cm(0.76)
        run = p.add_run(f"{i}. {title}: ")
        run.font.name = "Bookman Old Style"
        run.font.size = Pt(12)
        run.font.bold = True
        run = p.add_run(desc)
        run.font.name = "Bookman Old Style"
        run.font.size = Pt(12)


def create_chapter6(doc):
    """Chapter 6: Weekly Task."""
    add_page_break(doc)

    add_heading_styled(doc, "Chapter 6: Weekly Task", level=1)

    tasks = [
        ("Week 1-2", "Literature review on behavioral biometrics, keystroke dynamics, and "
         "touch-based authentication. Study of existing approaches and identification of research gaps."),
        ("Week 3-4", "Design of system architecture. Selection of technology stack (Kotlin, "
         "Jetpack Compose, Hilt, Room). Setup of Android development environment."),
        ("Week 5-6", "Development of core banking application: Login screen with custom PIN pad, "
         "Dashboard, and Transfer screens. Implementation of navigation using Jetpack Navigation."),
        ("Week 7-8", "Implementation of behavioral data collectors: KeystrokeCollector, "
         "TouchDataCollector, and SensorDataCollector. Integration with banking screens for "
         "transparent data capture."),
        ("Week 9-10", "Development of experiment management system: ExperimentViewModel, "
         "participant management, session types (Enrollment, Genuine, Impostor), and data export "
         "functionality."),
        ("Week 11-12", "Data collection from 8 real participants. Export and verification of "
         "collected CSV data (PIN keystrokes, touches, motion, metadata)."),
        ("Week 13-14", "Synthetic data generation for 32 additional participants with age "
         "demographics. Development of statistical profile augmentation pipeline."),
        ("Week 15-16", "ML model development: feature extraction (124 features), training of "
         "Random Forest, MLP, and One-Class SVM classifiers. Achievement of 92.5% accuracy."),
        ("Week 17-18", "Generation of publication-quality visualizations (8 figures). "
         "ROC curves, confusion matrices, t-SNE plots, feature importance rankings."),
        ("Week 19-20", "Report writing, IEEE paper preparation, and documentation. "
         "Final testing and verification of all components."),
    ]

    table = doc.add_table(rows=len(tasks) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for j, h in enumerate(["Week", "Task Description"]):
        cell = table.cell(0, j)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = "Bookman Old Style"
        run.font.size = Pt(11)
        run.font.bold = True

    for i, (week, task) in enumerate(tasks):
        for j, val in enumerate([week, task]):
            cell = table.cell(i + 1, j)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = "Bookman Old Style"
            run.font.size = Pt(11)

    add_formatted_paragraph(doc, "Table 3: Weekly Task Schedule",
        size=10, bold=True, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=4)


def create_references(doc):
    """References page."""
    add_page_break(doc)

    add_heading_styled(doc, "References", level=1)

    refs = [
        '[1] Statista, "Number of mobile banking users worldwide from 2020 to 2025," '
        'Statista Research Department, 2024.',

        '[2] H. Khan and U. Hengartner, "Towards application-centric implicit authentication '
        'on smartphones," in Proc. 15th Workshop on Mobile Computing Systems and Applications, '
        '2014, pp. 1-6.',

        '[3] P. S. Teh, A. B. J. Teoh, and S. Yue, "A survey of keystroke dynamics biometrics," '
        'The Scientific World Journal, vol. 2013, pp. 1-24, 2013.',

        '[4] F. Monrose and A. Rubin, "Authentication via keystroke dynamics," in Proc. 4th ACM '
        'Conference on Computer and Communications Security, 1997, pp. 48-56.',

        '[5] N. Zheng, K. Bai, H. Huang, and H. Wang, "You are how you touch: User verification '
        'on smartphones via tapping behaviors," in Proc. IEEE 22nd International Conference on '
        'Network Protocols, 2014, pp. 221-232.',

        '[6] A. Acien, A. Morales, J. Fierrez, R. Vera-Rodriguez, and O. Delgado-Mohatar, '
        '"TypeNet: Deep learning keystroke biometrics," IEEE Trans. Biometrics, Behavior, and '
        'Identity Science, vol. 4, no. 1, pp. 57-70, 2021.',

        '[7] M. Frank, R. Biedert, E. Ma, I. Martinovic, and D. Song, "Touchalytics: On the '
        'applicability of touchscreen input as a behavioral biometric for continuous authentication," '
        'IEEE Trans. Information Forensics and Security, vol. 8, no. 1, pp. 136-148, 2013.',

        '[8] A. Serwadda and V. V. Phoha, "Examining a large keystroke biometrics dataset for '
        'statistical-attack openings," ACM Trans. Information and System Security, vol. 16, no. 2, '
        'pp. 1-30, 2013.',

        '[9] C. Bo, L. Zhang, X.-Y. Li, Q. Huang, and Y. Wang, "SilentSense: Silent user '
        'identification via touch and movement behavioral biometrics," in Proc. 19th Annual '
        'International Conference on Mobile Computing and Networking, 2013, pp. 187-190.',

        '[10] Z. Sitova et al., "HMOG: New behavioral biometric features for continuous '
        'authentication of smartphone users," IEEE Trans. Information Forensics and Security, '
        'vol. 11, no. 5, pp. 877-892, 2015.',

        '[11] M. Temper, S. Tjoa, and M. Kaiser, "Multimodal behavioral biometrics for continuous '
        'mobile authentication," in Proc. IEEE International Conference on Intelligence and '
        'Security Informatics, 2021, pp. 1-6.',

        '[12] L. Findlater et al., "Age-related differences in performance with touchscreens '
        'compared to traditional mouse input," in Proc. SIGCHI Conference on Human Factors in '
        'Computing Systems, 2013, pp. 343-346.',
    ]

    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.hanging_indent = Cm(0.76)
        run = p.add_run(ref)
        run.font.name = "Bookman Old Style"
        run.font.size = Pt(11)


def main():
    print("=" * 60)
    print("  Generating Project Progress Report (DOCX)")
    print("=" * 60)

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Bookman Old Style"
    font.size = Pt(12)

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # Build report
    print("  📄 Creating Title Page...")
    create_title_page(doc)

    print("  📄 Creating Declaration...")
    create_declaration_page(doc)

    print("  📄 Creating Table of Contents...")
    create_toc(doc)

    print("  📄 Chapter 1: Introduction and Problem Statement...")
    create_chapter1(doc)

    print("  📄 Chapter 2: Objectives...")
    create_chapter2(doc)

    print("  📄 Chapter 3: Project Work Carried Out...")
    create_chapter3(doc)

    print("  📄 Chapter 4: Results and Analysis...")
    create_chapter4(doc)

    print("  📄 Chapter 5: Future Work Plan...")
    create_chapter5(doc)

    print("  📄 Chapter 6: Weekly Task...")
    create_chapter6(doc)

    print("  📄 References...")
    create_references(doc)

    # Save
    doc.save(OUTPUT_PATH)

    print(f"\n{'=' * 60}")
    print(f"  ✅ REPORT GENERATED!")
    print(f"  📄 Output: {OUTPUT_PATH}")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()
